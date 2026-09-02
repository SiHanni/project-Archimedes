"""
심사 제출용 DOCX 보고서 생성.

HTML 을 LibreOffice 로 변환하면 표 서식과 스타일이 깨진다(실측: 표 테두리·정렬이
모두 유실). 그래서 HTML 을 파싱해 **python-docx 로 문서를 직접 조립한다.**
제목 계층·표 스타일·이미지 배치를 워드 기본 스타일에 맞춰 넣으므로, 심사위원이
받아서 그대로 열람·주석·인쇄할 수 있다.

⚠️ 본문은 `~/Desktop/TIPS 기능 심사 final.html` 을 정본으로 삼는다. 문구를 고칠 때는
   HTML 을 고치고 이 스크립트를 다시 돌린다. 두 곳을 따로 고치면 반드시 어긋난다.
"""

from __future__ import annotations

import re
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
import json as _json

SRC = Path.home() / "Desktop" / "TIPS 기능 심사 final.html"
DST = Path.home() / "Desktop" / "TIPS 기능 심사 final.docx"
IMG = Path("/Users/sihwan/project-Archimedes/web/public/spec")

# 본문 글꼴. 한글이 깨지지 않도록 동아시아 글꼴을 함께 지정해야 한다.
FONT = "맑은 고딕"


class Block:
    def __init__(self, kind: str, **kw):
        self.kind = kind
        self.__dict__.update(kw)


class Parser(HTMLParser):
    """필요한 구조만 뽑는다 — 제목·문단·표·강조 박스."""

    def __init__(self) -> None:
        super().__init__()
        self.blocks: list[Block] = []
        self.stack: list[str] = []
        self.buf: list[str] = []
        self.row: list[str] = []
        self.table: list[list[str]] = []
        self.in_head = False
        self.note_kind = ""

    def handle_starttag(self, tag, attrs):
        a = dict(attrs)
        cls = a.get("class", "")
        if tag in ("h1", "h2", "h3", "h4", "p", "li", "figcaption"):
            self.buf = []
            self.stack.append(tag)
        elif tag == "div" and ("note" in cls or cls.startswith("eq")):
            # ⚠️ 강조 박스 본문은 <p> 로 감싸이지 않은 경우가 많다. 그대로 두면
            #    DOCX 에서 통째로 누락된다(실측: 51개 박스가 사라졌다).
            #    div 안의 맨 텍스트도 모으도록 버퍼를 연다.
            self.note_kind = "warn" if "warn" in cls else ("good" if "good" in cls else "note")
            self.blocks.append(Block("note_start", note=self.note_kind))
            self.buf = []
            self.stack.append("div")
        elif tag == "table":
            self.table = []
        elif tag == "tr":
            self.row = []
        elif tag in ("td", "th"):
            self.buf = []
            self.stack.append(tag)
            self.in_head = tag == "th"
        elif tag == "br":
            self.buf.append(" ")

    def handle_endtag(self, tag):
        txt = re.sub(r"\s+", " ", "".join(self.buf)).strip()
        if tag in ("h1", "h2", "h3", "h4"):
            if txt:
                self.blocks.append(Block("head", level=int(tag[1]), text=txt))
            self.buf = []
        elif tag in ("p", "li", "figcaption"):
            if txt:
                self.blocks.append(Block("li" if tag == "li" else "para", text=txt))
            self.buf = []
        elif tag in ("td", "th"):
            self.row.append(txt)
            self.buf = []
        elif tag == "tr":
            if self.row:
                self.table.append(self.row)
            self.row = []
        elif tag == "table":
            if self.table:
                self.blocks.append(Block("table", rows=self.table))
            self.table = []
        elif tag == "div":
            if self.note_kind:
                if self.stack and self.stack[-1] == "div":
                    self.stack.pop()
                if txt:
                    self.blocks.append(Block("para", text=txt))
                self.buf = []
                self.blocks.append(Block("note_end"))
                self.note_kind = ""

    def handle_data(self, data):
        if self.stack:
            self.buf.append(data)

    # ⚠️ 엔티티(&nbsp; &gt; 등)를 처리하지 않으면 수식 본문이 그 지점에서 끊겨
    #    조각만 남는다. 수식 박스가 통째로 누락되던 원인이었다.
    def handle_entityref(self, name):
        if self.stack:
            import html as _h
            self.buf.append(_h.unescape(f"&{name};"))

    def handle_charref(self, name):
        if self.stack:
            import html as _h
            self.buf.append(_h.unescape(f"&#{name};"))


def set_font(doc: Document) -> None:
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    for name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        h = doc.styles[name]
        h.font.name = FONT
        h.font.color.rgb = RGBColor(0x1A, 0x1C, 0x22)
        h.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def keep_with_next(para) -> None:
    """이 문단이 다음 내용과 떨어지지 않게 한다(제목이 페이지 끝에 홀로 남는 것 방지)."""
    para.paragraph_format.keep_with_next = True
    para.paragraph_format.keep_together = True


def no_split(table) -> None:
    """
    표가 페이지 경계에서 잘리지 않게 한다.

    워드는 기본적으로 행 중간에서 페이지를 넘긴다. 수치표가 그렇게 잘리면 헤더와
    값이 분리돼 읽을 수 없다. 행 단위 분할 금지(cantSplit)와 헤더 반복(tblHeader)을
    함께 건다.
    """
    for i, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        trPr.append(cant)
        if i == 0:
            th = OxmlElement("w:tblHeader")
            trPr.append(th)


def add_table(host, rows: list[list[str]]) -> None:
    """host 는 Document 또는 _Cell. 컨테이너 셀 안에도 표를 넣을 수 있어야 한다."""
    cols = max(len(r) for r in rows)
    t = host.add_table(rows=0, cols=cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, r in enumerate(rows):
        cells = t.add_row().cells
        for j in range(cols):
            v = r[j] if j < len(r) else ""
            para = cells[j].paragraphs[0]
            run = para.add_run(v)
            run.font.size = Pt(8.5)
            run.font.name = FONT
            run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
            if i == 0:
                run.bold = True
            # 숫자 칸은 오른쪽 정렬 — 자릿수를 눈으로 맞추기 위함
            if re.fullmatch(r"[\d,.\s%±~+\-]+p?", v or ""):
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    no_split(t)
    host.add_paragraph()


def main() -> None:
    p = Parser()
    html = SRC.read_text()
    html = re.sub(r"<style.*?</style>", "", html, flags=re.S)
    p.feed(html)

    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.0)
        s.left_margin = s.right_margin = Cm(1.8)
    set_font(doc)

    in_note = False
    for b in p.blocks:
        if b.kind == "head":
            if b.level == 1:
                h = doc.add_heading(b.text, 0)
            else:
                h = doc.add_heading(b.text, min(b.level - 1, 4))
            for r in h.runs:
                r.font.name = FONT
                r.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
            keep_with_next(h)
        elif b.kind == "para":
            para = doc.add_paragraph()
            run = para.add_run(b.text)
            run.font.name = FONT
            run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
            if in_note:
                para.paragraph_format.left_indent = Cm(0.5)
                run.font.size = Pt(9.5)
        elif b.kind == "li":
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(b.text)
            run.font.name = FONT
            run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        elif b.kind == "table":
            add_table(doc, b.rows)
        elif b.kind == "note_start":
            in_note = True
        elif b.kind == "note_end":
            in_note = False

    # ── 부록: 실사진 10건 이미지 ──
    doc.add_page_break()
    h = doc.add_heading("부록. 실거래 사진 10건 대조 결과", 1)
    for r in h.runs:
        r.font.name = FONT
        r.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    # 각 건의 픽셀·거리 수치를 이미지와 함께 싣는다. 시험 요구사항이 "픽셀 결과와
    # 거리가 나온 화면 캡처"이므로, 사진만 있고 수치가 없으면 요구를 충족하지 못한다.
    data = _json.loads((IMG / "outline.json").read_text())
    KOR = {"ring": "반지", "necklace": "목걸이", "earring": "귀걸이", "goldbar": "골드바"}

    order = ["T379_01", "T374_01", "T390_01", "T332_01", "T338_01",
             "T384_01", "T341_01", "T192_01", "T330_01", "T152_01"]
    for idx, k in enumerate(order, 1):
        files = [IMG / f"{k}_{sfx}.jpg" for sfx in ("orig", "overlay", "diff")]
        if not all(f.exists() for f in files):
            continue
        v = data.get(k, {})
        d = v.get("dist", {})
        item = v.get("item", "")

        cap = doc.add_paragraph()
        cr = cap.add_run(f"{idx}. {k.replace('_01', '')}  {item}")
        cr.bold = True
        cr.font.size = Pt(10)
        cr.font.name = FONT
        cr.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        keep_with_next(cap)

        # ⚠️ 한 항목 = **표 하나**로 만든다. 사진표·수치표를 따로 두면 그 사이에서
        #    페이지가 갈리고, 표 안에 표를 넣으면(중첩) 워드가 블록을 통째로 다음
        #    페이지로 밀어 아래 2/3 가 빈다. 단일 표 + 행 분할 금지가 가장 안정적이다.
        t = doc.add_table(rows=4, cols=3)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER

        for j, (f, label) in enumerate(zip(files, ("원본", "모델 추출 영역", "정답 대조"))):
            cp = t.cell(0, j).paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.add_run().add_picture(str(f), width=Cm(4.9))
            lp = t.cell(1, j).paragraphs[0]
            lr = lp.add_run(label)
            lr.font.size = Pt(8)
            lr.font.name = FONT
            lr.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
            lp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def fill(row: int, text: str) -> None:
            merged = t.cell(row, 0).merge(t.cell(row, 2))
            para = merged.paragraphs[0]
            run = para.add_run(text)
            run.font.size = Pt(8.5)
            run.font.name = FONT
            run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

        fill(2, "시험항목 2 (픽셀)   "
                f"정답 {v.get('gt_px', 0):,} px · 모델 {v.get('model_px', 0):,} px · "
                f"교집합 {v.get('inter_px', 0):,} px · 누락 {v.get('miss_px', 0):,} px · "
                f"과다 {v.get('over_px', 0):,} px · IoU {v.get('iou', 0):.2f}% · "
                f"재현율 {v.get('recall', 0):.1f}% · 정밀도 {v.get('precision', 0):.1f}%")
        fill(3, "시험항목 3 (거리, 참고)   "
                f"가정 품목 {KOR.get(d.get('k'), '')} · 추정 {d.get('cm', 0)} cm · "
                f"범위 {d.get('lo', 0)}~{d.get('hi', 0)} cm · 오차 ±{d.get('sig', 0)}% · "
                f"원본 {v.get('w', 0)}×{v.get('h', 0)} · 거래일 {v.get('date', '')}")
        no_split(t)
        doc.add_paragraph()

    # ── 부록 B: 관리자 화면 캡처 ──
    # 시험 요구사항: 태깅 검증과 매칭 점수를 "직접 제품 화면에서" 확인한 결과를 제출한다.
    shots = IMG / "shots"
    have = sorted(shots.glob("*.png")) if shots.exists() else []
    if have:
        doc.add_page_break()
        h = doc.add_heading("부록 B. 관리자 화면 검증 캡처", 1)
        for r in h.runs:
            r.font.name = FONT
            r.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        lp = shots / "labels.json"
        LABEL = _json.loads(lp.read_text()) if lp.exists() else {}
        pics = [f for f in have if f.suffix == ".png"]
        # 캡처 1장이 폭 17cm 기준 8~9cm 다. 페이지(본문 25.7cm)에 2~3장이 들어가므로
        # 장마다 페이지를 넘기지 않는다. 대신 캡션과 그림이 갈리지 않게만 묶는다.
        for f in pics:
            cap = doc.add_paragraph()
            cr = cap.add_run(LABEL.get(f.stem, f.stem))
            cr.bold = True
            cr.font.size = Pt(10)
            cr.font.name = FONT
            cr.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
            keep_with_next(cap)
            ip = doc.add_paragraph()
            ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ip.paragraph_format.keep_together = True
            ip.paragraph_format.space_after = Pt(14)
            ip.add_run().add_picture(str(f), width=Cm(17.0))

    doc.save(DST)
    print(f"생성: {DST}  ({DST.stat().st_size/1024:.0f} KB)")
    print(f"블록 {len(p.blocks)}개 · 표 {sum(1 for b in p.blocks if b.kind=='table')}개")


if __name__ == "__main__":
    main()
